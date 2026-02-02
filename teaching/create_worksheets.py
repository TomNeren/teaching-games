#!/usr/bin/env python3
"""
Erstellt Arbeitsblätter im Tom-Format für die Unterrichtseinheiten
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Hintergrundfarbe für Zelle setzen"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_header_table(doc, title, lernziele, fach="Politik/Geschichte", niveau="B2"):
    """Erstellt die Header-Tabelle im Tom-Format"""
    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    
    # Zeile 1: Titel und Metadaten
    cell = table.cell(0, 0)
    cell.text = f"Arbeitsblatt\n\n{title}"
    cell.paragraphs[0].runs[0].bold = True
    set_cell_shading(cell, "E6E6E6")
    
    # Lernziele
    cell = table.cell(0, 1)
    cell.text = f"Lernziele:\n\n{lernziele}"
    cell.merge(table.cell(0, 2))
    
    # Zeile 2: Fach und Niveau
    table.cell(1, 0).text = "Name:                    Vorname:"
    table.cell(1, 1).text = fach
    table.cell(1, 2).text = f"Niveau: {niveau}"
    
    # Datum-Zeile
    row = table.add_row()
    row.cells[0].text = ""
    row.cells[1].text = ""
    row.cells[2].text = "Datum:"
    
    doc.add_paragraph()

def create_checkbox_table(doc, title, items):
    """Erstellt eine Tabelle mit Checkboxen"""
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Header
    table.cell(0, 0).text = "Kriterium"
    table.cell(0, 1).text = "✓"
    table.cell(0, 2).text = "Notizen"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "D9D9D9")
    
    # Items
    for item in items:
        row = table.add_row()
        row.cells[0].text = item
        row.cells[1].text = "☐"
        row.cells[2].text = ""
    
    doc.add_paragraph()

def create_demokratie_worksheet():
    """Arbeitsblatt A: Demokratie → Autokratie"""
    doc = Document()
    
    # Header
    create_header_table(
        doc,
        "Von der Demokratie zur Autokratie\nDas Fünf-Stufen-Modell",
        "• Warnsignale autoritärer Entwicklungen erkennen\n• Historische Beispiele analysieren\n• Mechanismen des demokratischen Verfalls verstehen",
        "Politik/Geschichte",
        "B2"
    )
    
    # Einführung
    p = doc.add_paragraph()
    p.add_run("Das Fünf-Stufen-Modell des demokratischen Verfalls").bold = True
    doc.add_paragraph("Demokratien sterben heute meist nicht durch Putsche, sondern durch schrittweise Erosion von innen.")
    
    # Die fünf Stufen
    stufen = [
        ("🟢 Stufe 1", "Die verwundbare Demokratie", "Funktionierende Institutionen, aber wachsende Unzufriedenheit"),
        ("🟡 Stufe 2", "Der Aufstieg des 'Retters'", "Charismatischer Führer mit populistischer Rhetorik"),
        ("🟠 Stufe 3", "Eroberung der Institutionen", "Justiz, Medien, Wahlrecht werden übernommen"),
        ("🔴 Stufe 4", "Die gelenkte Demokratie", "Wahlen finden statt, sind aber nicht fair"),
        ("⚫ Stufe 5", "Die offene Autokratie", "Keine Opposition, totale Kontrolle"),
    ]
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.cell(0, 0).text = "Stufe"
    table.cell(0, 1).text = "Bezeichnung"
    table.cell(0, 2).text = "Kennzeichen"
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "D9D9D9")
    
    for stufe, name, beschreibung in stufen:
        row = table.add_row()
        row.cells[0].text = stufe
        row.cells[1].text = name
        row.cells[2].text = beschreibung
    
    doc.add_paragraph()
    
    # Warnsignale Checkliste
    create_checkbox_table(doc, "Die vier Warnsignale (nach Levitsky/Ziblatt)", [
        "Ablehnung demokratischer Spielregeln",
        "Delegitimierung politischer Gegner als 'Feinde'",
        "Tolerierung oder Ermutigung von Gewalt",
        "Bereitschaft, Grundrechte einzuschränken"
    ])
    
    # Aufgabe 1
    p = doc.add_paragraph()
    p.add_run("Aufgabe 1: Fallstudie analysieren").bold = True
    doc.add_paragraph("Wähle eines der folgenden Länder und ordne die Entwicklungen den fünf Stufen zu:")
    doc.add_paragraph("☐ Weimar → Drittes Reich (1919-1933)")
    doc.add_paragraph("☐ Ungarn unter Orbán (2010-heute)")  
    doc.add_paragraph("☐ Türkei unter Erdoğan (2002-heute)")
    doc.add_paragraph("☐ Venezuela (1999-heute)")
    
    doc.add_paragraph()
    doc.add_paragraph("Meine Analyse:")
    for i in range(6):
        doc.add_paragraph("_" * 80)
    
    # Aufgabe 2
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Aufgabe 2: Warnsignal-Check").bold = True
    doc.add_paragraph("Analysiere das folgende Zitat. Welche Warnsignale erkennst du?")
    doc.add_paragraph()
    
    # Zitatbox
    p = doc.add_paragraph()
    p.add_run('"Die Presse ist der Feind des Volkes. Nur ich sage euch die Wahrheit."').italic = True
    doc.add_paragraph()
    
    doc.add_paragraph("Erkannte Warnsignale: " + "_" * 50)
    doc.add_paragraph("Begründung: " + "_" * 60)
    for i in range(3):
        doc.add_paragraph("_" * 80)
    
    # Reflexion
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Reflexion").bold = True
    doc.add_paragraph("Was können Bürger tun, um demokratischen Verfall zu stoppen?")
    for i in range(4):
        doc.add_paragraph("_" * 80)
    
    doc.save('/root/.openclaw/workspace/teaching/AB_Demokratie_Autokratie.docx')
    print("✓ AB_Demokratie_Autokratie.docx erstellt")

def create_simulation_worksheet():
    """Arbeitsblatt B: Simulationsspiel Reflexion"""
    doc = Document()
    
    # Header
    create_header_table(
        doc,
        '"Paragraph und Papier"\nReflexionsbogen zum Simulationsspiel',
        "• Die Bürokratie der Ausgrenzung verstehen\n• Historische Empathie entwickeln\n• Mechanismen der NS-Diskriminierung analysieren",
        "Geschichte",
        "B2"
    )
    
    # Meine Rolle
    p = doc.add_paragraph()
    p.add_run("Meine Rolle im Spiel").bold = True
    
    doc.add_paragraph("Name meiner Figur: " + "_" * 40)
    doc.add_paragraph("Beruf: " + "_" * 50)
    doc.add_paragraph("Meine Klassifizierung: ☐ Deutschblütig  ☐ Mischling 2. Grades  ☐ Mischling 1. Grades  ☐ 'Jude'")
    
    doc.add_paragraph()
    
    # Gefühle
    p = doc.add_paragraph()
    p.add_run("Teil 1: Gefühle benennen").bold = True
    
    doc.add_paragraph("Als ich meine Klassifizierung erfuhr, fühlte ich...")
    for i in range(3):
        doc.add_paragraph("_" * 80)
    
    doc.add_paragraph()
    doc.add_paragraph("Am meisten hat mich überrascht/erschüttert, dass...")
    for i in range(3):
        doc.add_paragraph("_" * 80)
    
    doc.add_paragraph()
    
    # Analyse
    p = doc.add_paragraph()
    p.add_run("Teil 2: Das System verstehen").bold = True
    
    create_checkbox_table(doc, "Was machte das System so 'effektiv'?", [
        "Alles war 'legal' und durch Gesetze geregelt",
        "Bürokratie machte Ausgrenzung unpersönlich",
        "Schrittweise Eskalation — kein plötzlicher Bruch",
        "Viele 'normale' Menschen machten mit",
        "Betroffene wurden isoliert und konnten sich kaum wehren"
    ])
    
    doc.add_paragraph("Hannah Arendt sprach von der 'Banalität des Bösen'. Was bedeutet das?")
    for i in range(4):
        doc.add_paragraph("_" * 80)
    
    doc.add_paragraph()
    
    # Transfer
    p = doc.add_paragraph()
    p.add_run("Teil 3: Transfer").bold = True
    
    doc.add_paragraph("Gibt es heute Situationen, in denen Menschen kategorisiert oder ausgeschlossen werden?")
    doc.add_paragraph("☐ Ja  ☐ Nein  ☐ Weiß nicht")
    doc.add_paragraph()
    doc.add_paragraph("Wenn ja, welche? Was sind Gemeinsamkeiten und Unterschiede zu 1935?")
    for i in range(4):
        doc.add_paragraph("_" * 80)
    
    doc.add_paragraph()
    
    # Abschluss
    p = doc.add_paragraph()
    p.add_run("Was ich aus dieser Stunde mitnehme:").bold = True
    for i in range(4):
        doc.add_paragraph("_" * 80)
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Zitat zum Nachdenken:").italic = True
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('"Ausgrenzung beginnt nicht mit Lagern. Sie beginnt mit Listen."').italic = True
    
    doc.save('/root/.openclaw/workspace/teaching/AB_Simulation_Reflexion.docx')
    print("✓ AB_Simulation_Reflexion.docx erstellt")

def create_rollenkarten():
    """Rollenkarten für das Simulationsspiel"""
    doc = Document()
    
    rollen = [
        {
            "name": "Maria Hoffmann",
            "beruf": "Lehrerin, 28 Jahre",
            "familie": "Vater Protestant, Mutter Katholikin",
            "dokument1": "Alle Großeltern christlich getauft",
            "dokument2": "Großvater mütterlicherseits: konvertiert 1890 vom Judentum zum Protestantismus",
            "klassifizierung": "Mischling 2. Grades"
        },
        {
            "name": "Friedrich Weber",
            "beruf": "Arzt, 35 Jahre",
            "familie": "Lutherische Familie seit Generationen",
            "dokument1": "Keine Auffälligkeiten",
            "dokument2": "Alle Dokumente vollständig, 'arische' Abstammung bestätigt",
            "klassifizierung": "Deutschblütig"
        },
        {
            "name": "Eva Rosenthal",
            "beruf": "Schneiderin, 22 Jahre",
            "familie": "Vater jüdisch, Mutter evangelisch",
            "dokument1": "Wuchs in gemischter Nachbarschaft auf",
            "dokument2": "Vater und beide Großeltern väterlicherseits jüdisch",
            "klassifizierung": "Mischling 1. Grades"
        },
        {
            "name": "Heinrich Braun",
            "beruf": "Beamter, 40 Jahre",
            "familie": "Katholische Familie aus Bayern",
            "dokument1": "Keine Auffälligkeiten bekannt",
            "dokument2": "Alle Dokumente in Ordnung",
            "klassifizierung": "Deutschblütig"
        },
        {
            "name": "Ruth Goldstein",
            "beruf": "Ärztin, 32 Jahre",
            "familie": "Jüdische Familie, nicht religiös praktizierend",
            "dokument1": "Beide Eltern jüdischer Abstammung",
            "dokument2": "Alle vier Großeltern jüdisch",
            "klassifizierung": "Jude (Volljude)"
        },
        {
            "name": "Karl Schmidt",
            "beruf": "Buchhalter, 45 Jahre",
            "familie": "Evangelische Familie",
            "dokument1": "Familiendokumente unvollständig (Krieg 1870)",
            "dokument2": "Nach Recherche: Alle Großeltern christlich",
            "klassifizierung": "Deutschblütig"
        },
        {
            "name": "Else Neumann",
            "beruf": "Krankenschwester, 26 Jahre",
            "familie": "Mutter jüdisch (konvertiert), Vater evangelisch",
            "dokument1": "Getauft, christlich erzogen",
            "dokument2": "Mutter vor der Ehe konvertiert, Großeltern mütterlicherseits jüdisch",
            "klassifizierung": "Mischling 1. Grades"
        },
        {
            "name": "Wilhelm Lange",
            "beruf": "Handwerker, 38 Jahre",
            "familie": "Arbeiterfamilie, evangelisch",
            "dokument1": "Keine Besonderheiten",
            "dokument2": "Alle Großeltern christlich getauft",
            "klassifizierung": "Deutschblütig"
        },
        {
            "name": "Margarete Stern",
            "beruf": "Studentin, 20 Jahre",
            "familie": "Vater evangelisch, Mutter jüdisch (praktizierend)",
            "dokument1": "Besucht manchmal die Synagoge mit der Mutter",
            "dokument2": "Großeltern mütterlicherseits jüdisch, gehört jüdischer Gemeinde an",
            "klassifizierung": "Geltungsjude"
        },
        {
            "name": "Ernst Müller",
            "beruf": "Lehrer, 50 Jahre",
            "familie": "Alteingesessene protestantische Familie",
            "dokument1": "Familie seit 1650 in der Region nachweisbar",
            "dokument2": "Vollständiger Stammbaum, keine 'Auffälligkeiten'",
            "klassifizierung": "Deutschblütig"
        },
        {
            "name": "Anna Kohn",
            "beruf": "Verkäuferin, 24 Jahre",
            "familie": "Jüdische Familie, assimiliert",
            "dokument1": "Eltern nicht religiös, aber jüdischer Abstammung",
            "dokument2": "Alle vier Großeltern jüdisch (wenn auch nicht praktizierend)",
            "klassifizierung": "Jude (Volljude)"
        },
        {
            "name": "Georg Fischer",
            "beruf": "Student, 21 Jahre",
            "familie": "Katholische Familie",
            "dokument1": "Adoptiert als Säugling",
            "dokument2": "Leibliche Eltern unbekannt — Abstammung 'ungeklärt'",
            "klassifizierung": "Status ungeklärt (Sonderprüfung)"
        },
    ]
    
    for rolle in rollen:
        # Rollenkarte
        p = doc.add_paragraph()
        p.add_run("═" * 50).bold = True
        
        p = doc.add_paragraph()
        p.add_run(f"ROLLENKARTE: {rolle['name']}").bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(f"Beruf: {rolle['beruf']}")
        doc.add_paragraph(f"Familie: {rolle['familie']}")
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Dokument 1 (zu Beginn):").bold = True
        doc.add_paragraph(rolle['dokument1'])
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Dokument 2 (Station B):").bold = True
        doc.add_paragraph(rolle['dokument2'])
        
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Klassifizierung: ").bold = True
        p.add_run(rolle['klassifizierung'])
        
        p = doc.add_paragraph()
        p.add_run("═" * 50).bold = True
        
        doc.add_page_break()
    
    doc.save('/root/.openclaw/workspace/teaching/Rollenkarten_Simulation.docx')
    print("✓ Rollenkarten_Simulation.docx erstellt")

if __name__ == "__main__":
    create_demokratie_worksheet()
    create_simulation_worksheet()
    create_rollenkarten()
    print("\n✓ Alle Arbeitsblätter erstellt!")
