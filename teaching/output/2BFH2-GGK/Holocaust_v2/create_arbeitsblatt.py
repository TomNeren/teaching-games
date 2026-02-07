#!/usr/bin/env python3
"""
Arbeitsblatt für Holocaust-Stunde erstellen
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Vorlage laden
template_path = "/root/.openclaw/workspace/teaching/tom-skills-extracted/doc/arbeitsblatt-erstellen-v3/arbeitsblatt-erstellen-v3/templates/Vorlage_Fach.docx"
doc = Document(template_path)

# Header anpassen
for section in doc.sections:
    header = section.header
    for para in header.paragraphs:
        if "{{THEMA}}" in para.text:
            para.text = para.text.replace("{{THEMA}}", "Entscheidungen 1933-1945")
        if "{{FACH}}" in para.text:
            para.text = para.text.replace("{{FACH}}", "GGK")
        if "{{NIVEAU}}" in para.text:
            para.text = para.text.replace("{{NIVEAU}}", "C")
        if "{{DATUM}}" in para.text:
            para.text = para.text.replace("{{DATUM}}", "______")
    
    # Auch in Tabellen im Header
    for table in header.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.text = para.text.replace("{{THEMA}}", "Entscheidungen 1933-1945")
                    para.text = para.text.replace("{{FACH}}", "GGK")
                    para.text = para.text.replace("{{NIVEAU}}", "C")
                    para.text = para.text.replace("{{DATUM}}", "______")
                    para.text = para.text.replace("{{NAME}}", "Name: ___________________")

# Alle bestehenden Paragraphen löschen (außer Header/Footer)
for para in doc.paragraphs:
    p = para._element
    p.getparent().remove(p)

# Neuen Inhalt erstellen
def add_heading(doc, text, level=1):
    """Fügt Überschrift hinzu"""
    h = doc.add_heading(text, level=level)
    h.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_paragraph(doc, text, bold=False, italic=False):
    """Fügt Absatz hinzu"""
    p = doc.add_paragraph(text)
    if bold or italic:
        run = p.runs[0]
        run.font.bold = bold
        run.font.italic = italic
    return p

def add_form_line(doc, label, lines=1):
    """Fügt Formularzeile mit Linien hinzu"""
    p = doc.add_paragraph()
    p.add_run(label).font.bold = True
    p.add_run("\n")
    for i in range(lines):
        p.add_run("_" * 100 + "\n")
    return p

# --- SEITE 1: Self-Explanation-Prompts ---
add_heading(doc, "Teil 1: Meine Perspektive", level=1)

add_paragraph(doc, "Ich habe die Perspektive gewählt von:", bold=True)
p = doc.add_paragraph()
p.add_run("☐ David (jüdischer Junge, Opferperspektive)\n")
p.add_run("☐ Werner (deutscher Hitlerjunge, Zuschauerperspektive)")

doc.add_paragraph()  # Leerzeile

add_heading(doc, "Während du spielst: Entscheidungen reflektieren", level=2)

add_paragraph(doc, "Fülle das Arbeitsblatt aus, während du das Spiel spielst. "
              "Halte bei wichtigen Entscheidungen an und notiere deine Gedanken.", italic=True)

doc.add_paragraph()

# Entscheidung 1
add_heading(doc, "Entscheidung 1", level=3)
add_form_line(doc, "Was habe ich gewählt?", 1)
add_form_line(doc, "Warum habe ich mich so entschieden?", 2)
add_form_line(doc, "Was hätte passieren können, wenn ich anders gewählt hätte?", 2)
add_form_line(doc, "Was hat mein Charakter in diesem Moment gefühlt?", 2)

doc.add_paragraph()

# Entscheidung 2
add_heading(doc, "Entscheidung 2", level=3)
add_form_line(doc, "Was habe ich gewählt?", 1)
add_form_line(doc, "Warum habe ich mich so entschieden?", 2)
add_form_line(doc, "Was hätte passieren können, wenn ich anders gewählt hätte?", 2)
add_form_line(doc, "Was hat mein Charakter in diesem Moment gefühlt?", 2)

# Seitenumbruch
doc.add_page_break()

# --- SEITE 2: Perspektivwechsel ---
add_heading(doc, "Teil 2: Perspektivwechsel", level=1)

add_paragraph(doc, "Finde einen Partner/eine Partnerin, der/die die ANDERE Perspektive gespielt hat. "
              "Erzählt euch gegenseitig, was ihr erlebt habt.", bold=True)

doc.add_paragraph()

add_form_line(doc, "Mein Partner/Meine Partnerin hat die Perspektive von __________ gespielt.", 1)

doc.add_paragraph()

add_heading(doc, "Was mein Partner/meine Partnerin erlebt hat:", level=3)
add_form_line(doc, "Das ist ihm/ihr passiert:", 3)
add_form_line(doc, "Diese Entscheidungen musste er/sie treffen:", 3)
add_form_line(doc, "So hat sich sein/ihr Charakter gefühlt:", 3)

doc.add_paragraph()

add_heading(doc, "Meine Überraschung:", level=3)
add_form_line(doc, "Das hat mich überrascht, weil...", 3)

add_paragraph(doc, "💡 Tipp: Nutze Satzanfänge wie:", italic=True)
add_paragraph(doc, "• \"Ich habe mich für ... entschieden, weil...\"\n"
              "• \"Das hat mich überrascht, weil...\"\n"
              "• \"In dieser Situation hätte ich...\"\n"
              "• \"Mein Charakter fühlte sich ..., weil...\"", italic=True)

# Seitenumbruch
doc.add_page_break()

# --- SEITE 3: Reflexion ---
add_heading(doc, "Teil 3: Reflexion — Und ich?", level=1)

add_paragraph(doc, "Jetzt wird es persönlich. Denke über deine eigenen Erfahrungen nach.", bold=True)

doc.add_paragraph()

add_heading(doc, "Stilles Schreiben (10 Minuten)", level=3)

add_paragraph(doc, "Warum ist es so schwer, sich gegen die Mehrheit zu stellen?", bold=True)

# Große Schreibfläche
for i in range(12):
    doc.add_paragraph("_" * 100)

doc.add_paragraph()

add_heading(doc, "Gegenwartsbezug (optional):", level=3)
add_paragraph(doc, "Wo sehe ich heute ähnliche Mechanismen? (Beispiele: Schule, soziale Medien, Gesellschaft)")

for i in range(5):
    doc.add_paragraph("_" * 100)

doc.add_paragraph()
doc.add_paragraph()

add_heading(doc, "Exit-Ticket", level=3)
add_paragraph(doc, "Das Wichtigste, das ich heute gelernt habe...", bold=True)

for i in range(4):
    doc.add_paragraph("_" * 100)

# Footer anpassen (falls vorhanden)
for section in doc.sections:
    footer = section.footer
    for para in footer.paragraphs:
        if "{{" in para.text:
            para.text = para.text.replace("{{THEMA}}", "Entscheidungen 1933-1945")
            para.text = para.text.replace("{{DATUM}}", "2026")

# Speichern
output_path = "/root/.openclaw/workspace/teaching/output/2BFH2-GGK/Holocaust_v2/doc/AB_01_Entscheidungen.docx"
doc.save(output_path)
print(f"✓ Arbeitsblatt erstellt: {output_path}")
