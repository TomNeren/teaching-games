#!/usr/bin/env python3
"""
Arbeitsblatt/Klassenarbeit erstellen (v2)

Version: 2.0.1
Stand: 24.01.2025
"""

import argparse
import json
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH


# =============================================================================
# CONFIGURATION
# =============================================================================

SKILL_PATH = Path("/mnt/skills/user/arbeitsblatt-erstellen-v2")
TEMPLATE_AB = SKILL_PATH / "templates" / "Vorlage_Fach.docx"
TEMPLATE_KA = SKILL_PATH / "templates" / "Vorlage_Klassenarbeit.docx"

FALLBACK_PATHS = [
    Path("/mnt/user-data/uploads"),
    Path("/home/claude/arbeitsblatt-erstellen-v2/templates"),
]


# =============================================================================
# TEMPLATE SELECTION
# =============================================================================

def get_template(typ: str = "ab") -> Path:
    if typ.lower() in ["klassenarbeit", "ka", "test", "prüfung", "pruefung"]:
        template = TEMPLATE_KA
        filename = "Vorlage_Klassenarbeit.docx"
    else:
        template = TEMPLATE_AB
        filename = "Vorlage_Fach.docx"
    
    if template.exists():
        return template
    
    for fallback in FALLBACK_PATHS:
        path = fallback / filename
        if path.exists():
            print(f"ℹ️ Vorlage aus Fallback: {path}")
            return path
    
    raise FileNotFoundError(f"Vorlage nicht gefunden: {template}")


# =============================================================================
# HEADER REPLACEMENT
# =============================================================================

def replace_header_placeholders(doc: Document, replacements: dict) -> Document:
    """Ersetzt Platzhalter in der Header-Tabelle."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = "".join(run.text for run in paragraph.runs)
                    
                    modified = False
                    for old, new in replacements.items():
                        if old in full_text:
                            full_text = full_text.replace(old, new)
                            modified = True
                    
                    if modified and paragraph.runs:
                        paragraph.runs[0].text = full_text
                        for run in paragraph.runs[1:]:
                            run.text = ""
    
    return doc


# =============================================================================
# BODY REPLACEMENT
# =============================================================================

def clear_body(doc: Document) -> Document:
    """Entfernt alle Paragraphen nach der Header-Tabelle."""
    while len(doc.paragraphs) > 1:
        p = doc.paragraphs[-1]._element
        p.getparent().remove(p)
    
    if doc.paragraphs:
        doc.paragraphs[0].clear()
    
    return doc


def add_body_content(doc: Document, content: list) -> Document:
    """
    Fügt strukturierten Inhalt zum Body hinzu.
    
    Formatierung (passend zu Toms Vorlagen):
    - heading1: Heading 1 Style (Avenir Next Ultra Light, 16pt), unterstrichen
    - heading2: Heading 2 Style (Avenir Next Medium, 13pt), unterstrichen
    - paragraph: Normal Style (Avenir Next Regular)
    - list_item: Normal Style mit Nummerierung, eingerückt
    """
    list_counter = 0
    
    for item in content:
        item_type = item.get("type", "paragraph")
        text = item.get("text", "")
        
        if item_type == "heading1":
            # Hauptüberschrift: Heading 1 Style + unterstrichen
            list_counter = 0
            p = doc.add_paragraph(text, style='Heading 1')
            for run in p.runs:
                run.underline = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            
        elif item_type == "heading2":
            # Unterüberschrift: Heading 2 Style (Avenir Next Medium, 13pt), OHNE Unterstrich
            list_counter = 0
            p = doc.add_paragraph(text, style='Heading 2')
            # Kein Underline für heading2!
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            
        elif item_type == "paragraph":
            # Normaler Absatz: Normal Style
            p = doc.add_paragraph(text, style='Normal')
            if item.get("bold"):
                for run in p.runs:
                    run.bold = True
            p.paragraph_format.line_spacing = 1.3
            
        elif item_type == "list_item":
            # Nummerierte Liste: Normal Style mit manueller Nummer
            list_counter += 1
            p = doc.add_paragraph(style='Normal')
            p.add_run(f"({list_counter}) {text}")
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.line_spacing = 1.3
            
        elif item_type == "list_alpha":
            letter = item.get("letter", "a")
            p = doc.add_paragraph(style='Normal')
            p.add_run(f"{letter}) {text}")
            p.paragraph_format.left_indent = Inches(0.75)
            
        elif item_type == "empty":
            doc.add_paragraph()
            
        elif item_type == "box":
            p = doc.add_paragraph(style='Normal')
            run = p.add_run(f"💡 {text}")
            run.italic = True
            run.font.size = Pt(10)
    
    return doc


# =============================================================================
# VALIDATION
# =============================================================================

def validate_worksheet(doc_path: str) -> bool:
    """Validiert ein erstelltes Arbeitsblatt."""
    doc = Document(doc_path)
    errors = []
    
    all_text = ""
    for p in doc.paragraphs:
        all_text += p.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text + "\n"
    
    # Check 1: Nicht ersetzte Platzhalter
    placeholders = re.findall(r'\[[^\]]+\]', all_text)
    real_placeholders = [p for p in placeholders if p not in ['[]']]
    if real_placeholders:
        errors.append(f"❌ PLATZHALTER NICHT ERSETZT: {real_placeholders}")
    
    # Check 2: Falsche Umlaute
    umlaut_words = [
        'fuehren', 'koennen', 'muessen', 'waehlen', 'naechste',
        'aehnlich', 'ueber', 'fuer', 'Schueler', 'Pruefung', 
        'Uebung', 'erklaeren', 'begruenden', 'Ausfuehren'
    ]
    found_errors = [w for w in umlaut_words if w.lower() in all_text.lower()]
    if found_errors:
        errors.append(f"❌ UMLAUT-FEHLER: {found_errors}")
    
    if errors:
        print("\n🔴 VALIDIERUNG FEHLGESCHLAGEN:")
        print("\n".join(errors))
        return False
    
    print(f"✅ Validiert: {doc_path}")
    print(f"   Absätze: {len(doc.paragraphs)}, Tabellen: {len(doc.tables)}")
    return True


# =============================================================================
# MAIN WORKFLOW
# =============================================================================

def create_worksheet(config: dict, output_path: str, doc_type: str = "ab") -> str:
    template = get_template(doc_type)
    doc = Document(str(template))
    print(f"📄 Vorlage: {template.name}")
    
    replacements = {
        '[Thema]': config.get('thema', 'Thema'),
        '[Fach]': config.get('fach', 'Fach'),
        '[Ziel 1 – nicht länger als eine Zeile]': config.get('ziel1', ''),
        '[Ziel 2 – nicht länger als eine Zeile]': config.get('ziel2', ''),
        '[Ziel 3 – nicht länger als eine Zeile]': config.get('ziel3', ''),
        '[A / B / C]': config.get('niveau', 'B'),
    }
    
    if doc_type.lower() in ["ka", "klassenarbeit", "test"]:
        nummer = config.get('nummer', '1')
        replacements['Klassenarbeit [1/2/3/4]'] = f"Klassenarbeit {nummer}"
        replacements['[1/2/3/4]'] = nummer
        replacements['[abhängig von Aufbau]'] = str(config.get('punkte', ''))
    
    doc = replace_header_placeholders(doc, replacements)
    print("✓ Header ersetzt")
    
    content = config.get('content', [])
    if content:
        doc = clear_body(doc)
        doc = add_body_content(doc, content)
        print(f"✓ Body: {len(content)} Elemente")
    
    doc.save(output_path)
    print(f"💾 Gespeichert: {output_path}")
    
    validate_worksheet(output_path)
    
    return output_path


# =============================================================================
# CLI
# =============================================================================

def main():
    parser = argparse.ArgumentParser(description="Arbeitsblätter erstellen")
    parser.add_argument("--config", required=True, help="JSON-Datei oder JSON-String")
    parser.add_argument("--output", required=True, help="Ausgabepfad")
    parser.add_argument("--type", default="ab", choices=["ab", "ka"], help="ab oder ka")
    
    args = parser.parse_args()
    
    try:
        config = json.loads(args.config)
    except json.JSONDecodeError:
        config_path = Path(args.config)
        if config_path.exists():
            with open(config_path, encoding='utf-8') as f:
                config = json.load(f)
        else:
            raise FileNotFoundError(f"Config nicht gefunden: {args.config}")
    
    create_worksheet(config, args.output, args.type)


if __name__ == "__main__":
    main()
